from dataclasses import dataclass
from io import BytesIO
import os

import pandas as pd
import pypdfium2 as pdfium
import pytesseract
from docx import Document as DocxDocument
from PIL import Image
from pypdf import PdfReader

from app.core.config import get_settings
from app.db.models import Document
from app.services.storage import get_minio_client


class DocumentExtractionError(RuntimeError):
    """Raised when a local document cannot be safely extracted."""


@dataclass(frozen=True)
class ExtractedPage:
    page_number: int
    text: str
    source_name: str = ""
    source_type: str = ""


def extract_document(document: Document) -> list[ExtractedPage]:
    try:
        contents = _load_document_bytes(document)
        if document.source_type == "docx":
            extracted = _extract_docx(contents)
        elif document.source_type == "pdf":
            extracted = _extract_pdf(contents)
        elif document.source_type == "image":
            extracted = [(1, _ocr_image(Image.open(BytesIO(contents))))]
        elif document.source_type == "csv":
            _record_csv_metadata(document, contents)
            return []
        else:
            raise DocumentExtractionError("Unsupported document source type")
    except DocumentExtractionError:
        raise
    except Exception as error:
        raise DocumentExtractionError("Local document extraction failed") from error

    pages = [
        ExtractedPage(
            page_number=page_number,
            text=_normalize_text(text),
            source_name=document.original_filename,
            source_type=document.source_type,
        )
        for page_number, text in extracted
        if _normalize_text(text)
    ]
    if not pages:
        raise DocumentExtractionError("Document contains no extractable text")
    return pages


def _load_document_bytes(document: Document) -> bytes:
    response = get_minio_client().get_object(
        get_settings().minio_bucket,
        document.object_key,
    )
    try:
        return response.read()
    finally:
        response.close()
        response.release_conn()


def _extract_docx(contents: bytes) -> list[tuple[int, str]]:
    document = DocxDocument(BytesIO(contents))
    blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text.strip() for cell in row.cells))
    return [(1, "\n".join(blocks))]


def _extract_pdf(contents: bytes) -> list[tuple[int, str]]:
    reader = PdfReader(BytesIO(contents))
    pages: list[tuple[int, str]] = []
    rendered_document: pdfium.PdfDocument | None = None
    try:
        for page_index, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if not text.strip():
                if rendered_document is None:
                    rendered_document = pdfium.PdfDocument(contents)
                text = _ocr_pdf_page(rendered_document, page_index)
            pages.append((page_index + 1, text))
    finally:
        if rendered_document is not None:
            rendered_document.close()
    return pages


def _ocr_pdf_page(document: pdfium.PdfDocument, page_index: int) -> str:
    page = document[page_index]
    bitmap = page.render(scale=int(os.getenv("OCR_DPI", "300")) / 72)
    try:
        return _ocr_image(bitmap.to_pil())
    finally:
        bitmap.close()
        page.close()


def _ocr_image(image: Image.Image) -> str:
    return pytesseract.image_to_string(
        image,
        lang=os.getenv("OCR_LANGUAGES", "eng+vie"),
    )


def _record_csv_metadata(document: Document, contents: bytes) -> None:
    dataframe = pd.read_csv(BytesIO(contents))
    document.csv_schema = {
        str(column): str(dtype) for column, dtype in dataframe.dtypes.items()
    }
    document.csv_row_count = len(dataframe.index)


def _normalize_text(text: str) -> str:
    lines = [" ".join(line.split()) for line in text.replace("\x00", "").splitlines()]
    return "\n".join(line for line in lines if line).strip()
