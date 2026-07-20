from io import BytesIO
from uuid import uuid4

import pandas as pd
import pytest
from docx import Document as DocxDocument
from PIL import Image
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas

from app.core.config import Settings
from app.db.models import Document, DocumentStatus
from app.services import parsers
from app.services.parsers import extract_document


SETTINGS = Settings(database_url="sqlite://", jwt_secret="test-secret")


def _document(filename: str, source_type: str) -> Document:
    return Document(
        workspace_id=uuid4(),
        original_filename=filename,
        mime_type="application/octet-stream",
        object_key=f"workspace/{filename}",
        source_type=source_type,
        status=DocumentStatus.PROCESSING,
    )


def _docx_bytes() -> bytes:
    document = DocxDocument()
    document.add_paragraph("Quarterly policy")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Region"
    table.cell(0, 1).text = "Limit"
    table.cell(1, 0).text = "VN"
    table.cell(1, 1).text = "1250"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _text_pdf_bytes(text: str) -> bytes:
    output = BytesIO()
    pdf = canvas.Canvas(output)
    pdf.drawString(72, 720, text)
    pdf.showPage()
    pdf.save()
    return output.getvalue()


def _scanned_pdf_bytes() -> bytes:
    image = Image.new("RGB", (600, 200), "white")
    output = BytesIO()
    pdf = canvas.Canvas(output, pagesize=(600, 200))
    pdf.drawImage(ImageReader(image), 0, 0, width=600, height=200)
    pdf.showPage()
    pdf.save()
    return output.getvalue()


def test_docx_extracts_paragraphs_and_tables(monkeypatch: pytest.MonkeyPatch) -> None:
    document = _document("guide.docx", "docx")
    monkeypatch.setattr(parsers, "_load_document_bytes", lambda _: _docx_bytes())

    pages = extract_document(document, settings=SETTINGS)

    assert len(pages) == 1
    assert pages[0].page_number == 1
    assert pages[0].source_name == "guide.docx"
    assert pages[0].source_type == "docx"
    assert "Quarterly policy" in pages[0].text
    assert "Region | Limit" in pages[0].text
    assert "VN | 1250" in pages[0].text


def test_text_pdf_retains_page_number(monkeypatch: pytest.MonkeyPatch) -> None:
    document = _document("policy.pdf", "pdf")
    monkeypatch.setattr(
        parsers,
        "_load_document_bytes",
        lambda _: _text_pdf_bytes("Refunds are available within 30 days."),
    )

    pages = extract_document(document, settings=SETTINGS)

    assert [(page.page_number, page.source_name) for page in pages] == [
        (1, "policy.pdf")
    ]
    assert "Refunds are available within 30 days." in pages[0].text


def test_scanned_pdf_uses_configured_ocr_and_retains_page(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    document = _document("scan.pdf", "pdf")
    seen_calls: list[tuple[int, str]] = []
    monkeypatch.setattr(parsers, "_load_document_bytes", lambda _: _scanned_pdf_bytes())

    def fake_ocr(image: Image.Image, *, lang: str) -> str:
        seen_calls.append((image.width, lang))
        return "Scanned safety policy"

    monkeypatch.setattr(parsers.pytesseract, "image_to_string", fake_ocr)
    settings = Settings(
        database_url="sqlite://",
        jwt_secret="test-secret",
        ocr_languages="deu",
        ocr_dpi=144,
    )

    pages = extract_document(document, settings=settings)

    assert seen_calls == [(1200, "deu")]
    assert pages[0].page_number == 1
    assert pages[0].text == "Scanned safety policy"


def test_image_uses_local_ocr(monkeypatch: pytest.MonkeyPatch) -> None:
    document = _document("notice.png", "image")
    output = BytesIO()
    Image.new("RGB", (100, 40), "white").save(output, format="PNG")
    monkeypatch.setattr(parsers, "_load_document_bytes", lambda _: output.getvalue())
    monkeypatch.setattr(
        parsers,
        "_ocr_image",
        lambda _, *, languages: "Emergency notice",
    )

    pages = extract_document(document, settings=SETTINGS)

    assert pages[0].page_number == 1
    assert pages[0].text == "Emergency notice"
    assert pages[0].source_type == "image"


def test_csv_records_schema_and_row_count_without_prose_chunks(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    document = _document("sales.csv", "csv")
    contents = b"country,amount,approved\nVN,1250,true\nUS,500,false\n"
    monkeypatch.setattr(parsers, "_load_document_bytes", lambda _: contents)

    pages = extract_document(document, settings=SETTINGS)

    assert pages == []
    assert document.csv_row_count == 2
    assert document.csv_schema == {
        column: str(dtype)
        for column, dtype in pd.read_csv(BytesIO(contents)).dtypes.items()
    }
